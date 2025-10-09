import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from pathlib import Path
import os

class PdfToImageConverter:
    """
    PDF转Word工具类
    属性：
        pdf_path (Path): 输入PDF文件路径
        output_dir (Path): 输出文件夹路径（默认为PDF所在目录）
        image_dir (Path): 图片临时存储目录（默认output_dir/pdf_pages）
        docx_path (Path): 最终Word文档路径（默认output_dir/output_document.docx）
    
    方法：
        convert(): 执行完整转换流程
    """

    def __init__(
        self,
        pdf_path: str | Path,
        output_dir: str | Path = None,
        image_quality: int = 2  # 图片缩放比例（Matrix(2,2)=200%质量）
    ):
        """
        初始化转换器
        
        参数：
            pdf_path: 输入PDF文件路径
            output_dir: 输出文件夹路径（默认使用PDF所在目录）
            image_quality: 图片生成质量（建议1-3，数字越大清晰度越高但文件越大）
        """
        self.pdf_path = Path(pdf_path).resolve()
        self.output_dir = Path(output_dir or self.pdf_path.parent).resolve()
        self.image_dir = self.output_dir / "pdf_pages"
        self.docx_path = self.output_dir / "output_document.docx"
        self.image_quality = image_quality
        
        # 创建必要目录
        self.image_dir.mkdir(exist_ok=True, parents=True)
        self.output_dir.mkdir(exist_ok=True, parents=True)
        
        # 验证PDF文件存在
        if not self.pdf_path.is_file():
            raise FileNotFoundError(f"PDF文件未找到: {self.pdf_path}")

    def _pdf_to_images(self) -> list[Path]:
        """将PDF转换为高质量图片"""
        print(f"🔍 正在处理PDF: {self.pdf_path.name}")
        doc = fitz.open(str(self.pdf_path))
        image_paths = []
        
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            # 生成高清图片（调整matrix参数控制清晰度）
            matrix = fitz.Matrix(self.image_quality, self.image_quality)
            pix = page.get_pixmap(matrix=matrix)
            
            # 构建输出路径
            image_name = f"page_{page_num+1}.png"
            image_path = self.image_dir / image_name
            pix.save(str(image_path))
            
            image_paths.append(image_path)
            print(f"✅ 生成图片: {image_name}")
        
        doc.close()
        print(f"📂 已保存 {len(image_paths)} 张图片到: {self.image_dir}")
        return image_paths

    def _images_to_word(self, image_paths: list[Path]):
        """将图片插入Word文档"""
        print("\n📝 正在生成Word文档...")
        doc = Document()
        doc.add_heading('PDF转Word转换结果', 0)
        
        for idx, img_path in enumerate(image_paths, start=1):
            doc.add_heading(f'第 {idx} 页', level=1)
            doc.add_picture(str(img_path), width=Inches(6))
            doc.add_paragraph()
        
        doc.save(str(self.docx_path))
        print(f"✅ Word文档已保存: {self.docx_path}")

    def convert(self):
        """执行完整转换流程"""
        try:
            image_paths = self._pdf_to_images()
            self._images_to_word(image_paths)
            print("\n🎉 转换完成！")
            
        except Exception as e:
            print(f"❌ 发生错误: {str(e)}")

'''
# 使用示例
if __name__ == "__main__":
    # 配置参数
    pdf_file = Path("D:/PROJECT/Python/Convert/Catalog_beta_mini_2025.pdf")
    output_folder = Path("D:/OUTPUT")
    
    # 创建转换器实例
    converter = PdfConverter(
        pdf_path=pdf_file,
        output_dir=output_folder,
        image_quality=3  # 更高的质量（推荐2-3）
    )
    
    # 执行转换
    converter.convert()
'''